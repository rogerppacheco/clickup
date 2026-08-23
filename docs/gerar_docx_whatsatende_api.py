"""Gera Word com a documentação completa da API WhatsAtende (extraída do painel)."""
from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BASE_URL = "https://api.app14.whatsatende.com.br"
JS_PATH = Path(os.environ.get("TEMP", "/tmp")) / "wa_messages_api.js"
OUT_PATH = Path(r"C:\site-clickup\docs\WhatsAtende_API_Documentacao.docx")


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_endpoint(
    doc: Document,
    name: str,
    method: str,
    path: str,
    headers: list[tuple[str, str]] | None = None,
    payload: str | None = None,
    query_params: list[tuple[str, str]] | None = None,
) -> None:
    _add_heading(doc, name, level=3)
    _add_para(doc, f"Método: {method}", bold=True)
    _add_para(doc, "URL:")
    _add_code(doc, f"{BASE_URL}{path}")
    if headers:
        _add_para(doc, "Headers:", bold=True)
        for k, v in headers:
            _add_code(doc, f"{k}: {v}")
    if query_params:
        _add_para(doc, "Query params:", bold=True)
        for k, v in query_params:
            _add_para(doc, f"• {k}: {v}")
    if payload:
        _add_para(doc, "Body (JSON):", bold=True)
        _add_code(doc, payload.strip())
    doc.add_paragraph()


def parse_catalog_from_js(js: str) -> list[dict]:
    """Extrai endpoints do bundle MessagesAPI do frontend WhatsAtende."""
    idx = js.find('name:"Client"')
    if idx < 0:
        return []
    chunk = js[idx : idx + 25000]

    endpoints: list[dict] = []
    # name, method, url, optional payload, optional queryParams
    pattern = re.compile(
        r'name:"([^"]+)",method:"(GET|POST|PUT|DELETE|PATCH)",'
        r"url:`\$\{n\}([^`]+)`"
        r"(?:,headers:\[(.*?)\])?"
        r"(?:,payload:`([\s\S]*?)`)?"
        r"(?:,queryParams:\[(.*?)\])?",
        re.MULTILINE,
    )
    for m in pattern.finditer(chunk):
        name, method, path, headers_raw, payload, qparams_raw = m.groups()
        headers: list[tuple[str, str]] = []
        if headers_raw:
            for hm in re.finditer(
                r'key:"([^"]+)",value:"([^"]*)"', headers_raw
            ):
                headers.append((hm.group(1), hm.group(2)))
        query_params: list[tuple[str, str]] = []
        if qparams_raw:
            for qm in re.finditer(
                r'name:"([^"]+)",description:"([^"]*)"', qparams_raw
            ):
                query_params.append((qm.group(1), qm.group(2)))
        endpoints.append(
            {
                "name": name,
                "method": method,
                "path": path,
                "headers": headers,
                "payload": payload,
                "query_params": query_params,
            }
        )
    return endpoints


def build_fallback_endpoints() -> list[dict]:
    """Catálogo garantido (doc do painel + seção Contacts)."""
    bearer_conn = [("Authorization", "Bearer {{key_conexao}}"), ("Content-Type", "application/json")]
    bearer_token = [("Authorization", "Bearer {{token}}"), ("Content-Type", "application/json")]
    bearer_partner = [("Authorization", "Bearer {{key_partner}}"), ("Content-Type", "application/json")]

    return [
        {
            "section": "Client — Connections",
            "name": "Create Connection",
            "method": "POST",
            "path": "/api/messages/createConnection",
            "headers": bearer_token,
            "payload": '{\n  "companyToken": "{{companyToken}}",\n  "name": "Conexão WhatsApp1235"\n}',
        },
        {
            "section": "Client — Connections",
            "name": "Start Session",
            "method": "POST",
            "path": "/api/messages/startSession/{{whatsappId}}",
            "headers": bearer_conn,
        },
        {
            "section": "Client — Connections",
            "name": "Get QrCode",
            "method": "POST",
            "path": "/api/messages/getQrCode/{{whatsappId}}",
            "headers": bearer_conn,
        },
        {
            "section": "Client — Connections",
            "name": "Check Status",
            "method": "GET",
            "path": "/api/messages/checkStatus/{{whatsappId}}",
            "headers": bearer_conn,
        },
        {
            "section": "Client — Connections",
            "name": "Disconnect",
            "method": "POST",
            "path": "/api/messages/disconnect/{{whatsappId}}",
            "headers": bearer_conn,
        },
        {
            "section": "Client — Send",
            "name": "Send - Text",
            "method": "POST",
            "path": "/api/messages/send",
            "headers": bearer_conn,
            "payload": '{\n    "number": "{{Number}}",\n    "body": "Enviado via api",\n    "closeTicket": false\n}',
        },
        {
            "section": "Client — Send",
            "name": "Send - Image URL",
            "method": "POST",
            "path": "/api/messages/send/linkImage",
            "headers": bearer_conn,
            "payload": '{\n  "number": "{{Number}}",\n  "msdelay": 1000,\n  "url": "URL",\n  "caption": "Aqui está a imagem solicitada"\n}',
        },
        {
            "section": "Client — Send",
            "name": "Send - PDF URL",
            "method": "POST",
            "path": "/api/messages/send/linkPDF",
            "headers": bearer_conn,
            "payload": '{\n  "number": "{{Number}}",\n  "msdelay": 1000,\n  "fileUrl": "URL",\n  "caption": "Teste de envio de mídia pela API por link - Arquivo"\n}',
        },
        {
            "section": "Client — Send",
            "name": "Send - File Base64",
            "method": "POST",
            "path": "/api/messages/send/base64",
            "headers": bearer_conn,
            "payload": '{\n  "number": "{{Number}}",\n  "base64Data": "data:image/png;base64,...",\n  "fileName": "image.png",\n  "caption": "Esta é uma imagem via base64"\n}',
        },
        {
            "section": "Client — Send",
            "name": "Send - Media (multipart)",
            "method": "POST",
            "path": "/api/messages/sendMedia",
            "headers": [
                ("Authorization", "Bearer {{key_conexao}}"),
                ("Content-Type", "multipart/form-data"),
            ],
            "payload": "Campos (form-data): token, number, body, medias (arquivo).",
        },
        {
            "section": "Client — Contacts",
            "name": "Contact - Create",
            "method": "POST",
            "path": "/api/messages/contacts",
            "headers": bearer_conn,
            "payload": '{\n    "name": "{{Name}}",\n    "number": "{{Number}}",\n    "tagsIds": "1,2,3"\n}',
        },
        {
            "section": "Client — Contacts",
            "name": "Contact - Update",
            "method": "PUT",
            "path": "/api/messages/contacts/4",
            "headers": bearer_conn,
            "payload": '{\n    "name": "{{Name}}",\n    "number": "{{Number}}",\n    "tagsIds": "1,2,3"\n}',
        },
        {
            "section": "Client — Contacts",
            "name": "Contact - Delete",
            "method": "DELETE",
            "path": "/api/messages/contacts/4",
            "headers": bearer_conn,
        },
        {
            "section": "Client — Contacts",
            "name": "Contact - Search",
            "method": "GET",
            "path": "/api/messages/contacts/search?page=1&name={{Name}}&number={{Number}}",
            "headers": bearer_conn,
            "query_params": [
                ("page", "Número da página para paginação"),
                ("name", "Nome do contato a ser buscado"),
                ("number", "Número do telefone do contato (DDI+DDD+Número)"),
            ],
        },
        {
            "section": "Client — Contacts",
            "name": "Check Number",
            "method": "POST",
            "path": "/api/messages/checkNumber",
            "headers": bearer_conn,
            "payload": '{\n  "number": "{{Number}}"\n}',
        },
        {
            "section": "Client — Tickets",
            "name": "Ticket - Search",
            "method": "GET",
            "path": "/api/messages/tickets/search?ticketId={{ticketId}}&contactId={{contactId}}&page=1&status={{status}}",
            "headers": bearer_conn,
            "query_params": [
                ("page", "Número da página para paginação"),
                ("ticketId", "ID do ticket"),
                ("contactId", "ID do contato"),
                ("status", "Status do ticket (open, closed, pending, group, nps, lgpd)"),
            ],
        },
        {
            "section": "Partner — Plans",
            "name": "Create - Plans",
            "method": "POST",
            "path": "/api/plans",
            "headers": bearer_partner,
            "payload": '{\n  "name": "Plano 12",\n  "users": 10,\n  "connections": 10,\n  "queues": 10,\n  "amount": "100",\n  "useWhatsapp": true,\n  "useFacebook": true,\n  "useInstagram": true,\n  "useCampaigns": true,\n  "useSchedules": true,\n  "useInternalChat": true,\n  "useExternalApi": true,\n  "useKanban": true,\n  "trial": false,\n  "trialDays": 0,\n  "useOpenAi": true,\n  "useIntegrations": true\n}',
        },
        {
            "section": "Partner — Plans",
            "name": "Update - Plans",
            "method": "PUT",
            "path": "/api/plans/{{id_plan}}",
            "headers": bearer_partner,
            "payload": '(mesmo corpo do Create)',
        },
        {
            "section": "Partner — Plans",
            "name": "Delete - Plans",
            "method": "DELETE",
            "path": "/api/plans/{{id_plan}}",
            "headers": bearer_partner,
        },
        {
            "section": "Partner — Plans",
            "name": "Plans - All",
            "method": "GET",
            "path": "/plans",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Plans",
            "name": "Plans - ID",
            "method": "GET",
            "path": "/api/plans/1",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Company",
            "name": "Create - Company",
            "method": "POST",
            "path": "/api/companies",
            "headers": bearer_partner,
            "payload": '{\n  "name": "mkthub1",\n  "phone": "17991565280",\n  "email": "teste@teste.com.br",\n  "document": "04557950000172",\n  "planId": 1,\n  "partnerId": "1",\n  "password": "123456",\n  "dueDate": "2023-01-10T14:10:46.000Z"\n}',
        },
        {
            "section": "Partner — Company",
            "name": "Update - Company",
            "method": "PUT",
            "path": "/api/companies/{{companyId}}",
            "headers": bearer_partner,
            "payload": '(mesmo corpo do Create)',
        },
        {
            "section": "Partner — Company",
            "name": "Delete - Company",
            "method": "DELETE",
            "path": "/api/companies/{{companyId}}",
            "headers": bearer_partner,
        },
        {
            "section": "Partner — Company",
            "name": "Company All",
            "method": "GET",
            "path": "/api/companies",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Company",
            "name": "Company ID",
            "method": "GET",
            "path": "/api/companies/{{id_company}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Company",
            "name": "Company Email",
            "method": "GET",
            "path": "/api/companiesEmail/{{companiesEmail}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Users",
            "name": "Users Email",
            "method": "GET",
            "path": "/api/users/{{usersEmail}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Helper",
            "name": "Helper - Create",
            "method": "POST",
            "path": "/api/helps",
            "headers": bearer_partner,
            "payload": '{\n  "title": "teste",\n  "description": "teste de ajuda",\n  "video": "",\n  "link": ""\n}',
        },
        {
            "section": "Partner — Helper",
            "name": "Helper - Update",
            "method": "PUT",
            "path": "/api/helps/{{id_helps}}",
            "headers": bearer_partner,
            "payload": '(mesmo corpo do Create)',
        },
        {
            "section": "Partner — Helper",
            "name": "Helper - Delete",
            "method": "DELETE",
            "path": "/api/helps/{{id_helps}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Helper",
            "name": "Helper - All",
            "method": "GET",
            "path": "/api/helps",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Helper",
            "name": "Helper - ID",
            "method": "GET",
            "path": "/api/helps/{{id_helps}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Invoices",
            "name": "Create Invoice",
            "method": "POST",
            "path": "/api/invoices",
            "headers": bearer_partner,
            "payload": '{\n  "companyId": "1",\n  "amount": "100.00",\n  "status": "pending"\n}',
        },
        {
            "section": "Partner — Invoices",
            "name": "Get Invoice by ID",
            "method": "GET",
            "path": "/api/invoices/{{invoice_id}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
        {
            "section": "Partner — Invoices",
            "name": "Update Invoice",
            "method": "PUT",
            "path": "/api/invoices/{{invoice_id}}",
            "headers": bearer_partner,
            "payload": '{\n  "amount": "150.00",\n  "status": "paid"\n}',
        },
        {
            "section": "Partner — Invoices",
            "name": "Delete Invoice",
            "method": "DELETE",
            "path": "/api/invoices/{{invoice_id}}",
            "headers": [("Authorization", "Bearer {{key_partner}}")],
        },
    ]


def main() -> None:
    endpoints = build_fallback_endpoints()

    # Se o JS do painel existir, mescla payloads reais quando disponíveis
    if JS_PATH.exists():
        parsed = parse_catalog_from_js(JS_PATH.read_text(encoding="utf-8", errors="replace"))
        by_key = {(e["method"], e["path"].split("?")[0]): e for e in parsed}
        for ep in endpoints:
            key = (ep["method"], ep["path"].split("?")[0])
            # match flexível por path template
            for (m, p), src in by_key.items():
                if m == ep["method"] and (
                    p == ep["path"].split("?")[0]
                    or p.replace("{{whatsappId}}", "{{whatsappId}}") == ep["path"].split("?")[0]
                ):
                    if src.get("payload") and (
                        not ep.get("payload") or "data:image" in (ep.get("payload") or "")
                    ):
                        # mantém payload curto no base64
                        pass
                    if src.get("headers") and not ep.get("headers"):
                        ep["headers"] = src["headers"]
                    if src.get("query_params") and not ep.get("query_params"):
                        ep["query_params"] = src["query_params"]
                    break

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("WhatsAtende — Documentação da API", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _add_para(
        doc,
        "Fonte: painel WhatsAtende (app14) → menu API → Documentação API "
        "(/messages-api). Instância: api.app14.whatsatende.com.br. "
        "Versão observada da plataforma: 1.1.7d.",
    )
    _add_para(
        doc,
        "Documento gerado para a CLICKUP LTDA a partir da documentação "
        "embutida no frontend (SouChat/Whaticket) e da página Documentação API.",
        bold=False,
    )

    _add_heading(doc, "1. Visão geral", level=1)
    _add_para(doc, f"Base URL de produção: {BASE_URL}")
    _add_para(doc, "Autenticação:")
    _add_para(doc, "• key_conexao — Bearer token da conexão WhatsApp (menu Conexões → editar → token).")
    _add_para(doc, "• companyToken / token — usados na criação de conexão.")
    _add_para(doc, "• key_partner — token de parceiro (APIs Partner: planos, empresas, faturas).")
    _add_para(
        doc,
        "Número do destinatário: apenas dígitos, com DDI + DDD + número "
        "(ex.: 5531999999999), sem máscara.",
    )

    _add_heading(doc, "2. Observações importantes (painel)", level=1)
    _add_para(
        doc,
        "Antes de enviar mensagens, é necessário o cadastro do token vinculado "
        "à conexão que enviará as mensagens. Acesse o menu Conexões, clique em "
        "editar na conexão e insira o token no campo correspondente.",
    )
    _add_para(
        doc,
        "Esta documentação pública cobre envio (texto/mídia), conexão "
        "(QR/status), contatos e tickets. Webhooks de mensagem recebida / "
        "status, botões reply e lista de opções não aparecem nesta página "
        "e devem ser confirmados com o suporte (escopo técnico).",
    )

    _add_heading(doc, "3. Endpoints", level=1)

    current_section = None
    for ep in endpoints:
        section = ep.get("section") or "Geral"
        if section != current_section:
            current_section = section
            _add_heading(doc, section, level=2)
        _add_endpoint(
            doc,
            name=ep["name"],
            method=ep["method"],
            path=ep["path"],
            headers=ep.get("headers"),
            payload=ep.get("payload"),
            query_params=ep.get("query_params"),
        )

    _add_heading(doc, "4. Variáveis de integração (site-clickup)", level=1)
    _add_para(doc, "WHATSATENDE_API_URL — padrão: https://api.app14.whatsatende.com.br")
    _add_para(doc, "WHATSATENDE_TOKEN — Bearer da conexão (key_conexao)")
    _add_para(doc, "WHATSATENDE_WHATSAPP_ID — ID da conexão (whatsappId)")
    _add_para(doc, "Webhook inbound do bot: POST /api/crm/webhook-whatsapp/")

    _add_heading(doc, "5. Índice rápido de URLs", level=1)
    for ep in endpoints:
        _add_code(doc, f"{ep['method']:6} {BASE_URL}{ep['path']}  — {ep['name']}")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK: {OUT_PATH}")
    print(f"Endpoints: {len(endpoints)}")


if __name__ == "__main__":
    main()
